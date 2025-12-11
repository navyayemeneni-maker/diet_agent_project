"""
File Reader Module
==================
Handles reading medical reports from various file formats:
- PDF files (.pdf)
- Word documents (.docx)
- Text files (.txt)
- Images with OCR (.jpg, .png)

Author: Navya
Date: December 4, 2025
"""

import os
import PyPDF2
from docx import Document
from PIL import Image
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


class FileReader:
    """
    A utility class for reading text from various file formats.
    """
    
    def __init__(self):
        """Initialize the file reader"""
        self.supported_formats = {
            'pdf': 'PDF Document',
            'docx': 'Word Document',
            'doc': 'Word Document',
            'txt': 'Text File',
            'jpg': 'Image (requires OCR)',
            'jpeg': 'Image (requires OCR)',
            'png': 'Image (requires OCR)'
        }
        
        print("✅ File Reader initialized")
        print(f"📁 Supported formats: {', '.join(self.supported_formats.keys())}")
    
    
    def read_file(self, file_path):
        """
        Read text from a file based on its extension.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text from the file
        """
        
        # Check if file exists
        if not os.path.exists(file_path):
            return f"❌ Error: File not found at {file_path}"
        
        # Get file extension
        file_extension = file_path.lower().split('.')[-1]
        
        # Check if format is supported
        if file_extension not in self.supported_formats:
            return f"❌ Error: Unsupported file format '.{file_extension}'"
        
        print(f"\n📄 Reading file: {os.path.basename(file_path)}")
        print(f"📋 Format: {self.supported_formats[file_extension]}")
        
        try:
            # Route to appropriate reader
            if file_extension == 'pdf':
                text = self._read_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                text = self._read_word(file_path)
            elif file_extension == 'txt':
                text = self._read_text(file_path)
            elif file_extension in ['jpg', 'jpeg', 'png']:
                text = self._read_image(file_path)
            else:
                text = "❌ Error: Unsupported file format"
            
            if text and not text.startswith("❌"):
                print(f"✅ Successfully read {len(text)} characters")
                return text
            else:
                return text
                
        except Exception as e:
            error_msg = f"❌ Error reading file: {str(e)}"
            print(error_msg)
            return error_msg
    
    
    def _read_pdf(self, file_path):
        """Read text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get number of pages
                num_pages = len(pdf_reader.pages)
                print(f"   📖 PDF has {num_pages} page(s)")
                
                # Extract text from all pages
                text = ""
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                
                if not text.strip():
                    return "❌ Error: PDF appears to be empty or contains only images"
                
                return text.strip()
                
        except Exception as e:
            return f"❌ Error reading PDF: {str(e)}"
    
    
    def _read_word(self, file_path):
        """Read text from Word document"""
        try:
            doc = Document(file_path)
            
            # Get number of paragraphs
            num_paragraphs = len(doc.paragraphs)
            print(f"   📝 Document has {num_paragraphs} paragraph(s)")
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            if not text.strip():
                return "❌ Error: Word document appears to be empty"
            
            return text.strip()
            
        except Exception as e:
            return f"❌ Error reading Word document: {str(e)}"
    
    
    def _read_text(self, file_path):
        """Read text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                return "❌ Error: Text file is empty"
            
            return text.strip()
            
        except Exception as e:
            return f"❌ Error reading text file: {str(e)}"
    
    
    def _read_image(self, file_path):
        """Read text from image using OCR"""
        
        if not TESSERACT_AVAILABLE:
            return """❌ Error: OCR not available. 
            
To enable image reading, install Tesseract OCR:
- Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki
- Mac: brew install tesseract
- Linux: sudo apt-get install tesseract-ocr

Then try again."""
        
        try:
            # Open image
            image = Image.open(file_path)
            print(f"   🖼️ Image size: {image.size}")
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return "❌ Error: No text found in image. Image may not contain text or text is unclear."
            
            return text.strip()
            
        except Exception as e:
            return f"❌ Error reading image: {str(e)}"
    
    
    def get_supported_formats(self):
        """Return list of supported file formats"""
        return list(self.supported_formats.keys())
    
    
    def is_supported_format(self, file_path):
        """Check if file format is supported"""
        file_extension = file_path.lower().split('.')[-1]
        return file_extension in self.supported_formats


# ============================================================
# DEMO/TEST CODE
# ============================================================

def demo_file_reader():
    """
    Demonstration of the file reader capabilities
    """
    
    print("=" * 70)
    print("📁 FILE READER - DEMO")
    print("=" * 70)
    
    # Initialize reader
    reader = FileReader()
    
    print("\n" + "=" * 70)
    print("ℹ️  USAGE INSTRUCTIONS")
    print("=" * 70)
    
    print("""
To test the file reader, you can:

1. Create a test text file:
   - Open Notepad
   - Type some medical text
   - Save as 'test_report.txt' in your project folder

2. Or use an existing PDF/Word document

3. Then run this with the file path
    """)
    
    # Ask for file path
    print("\n" + "=" * 70)
    print("📂 TEST FILE READING")
    print("=" * 70)
    
    file_path = input("\nEnter file path (or press Enter to skip): ").strip()
    
    if file_path:
        # Remove quotes if user copied path with quotes
        file_path = file_path.strip('"').strip("'")
        
        # Read the file
        text = reader.read_file(file_path)
        
        print("\n" + "=" * 70)
        print("📄 EXTRACTED TEXT:")
        print("=" * 70)
        print(text)
        print("=" * 70)
    else:
        print("\n✅ Demo complete! File reader is ready to use.")
    
    print("\n" + "=" * 70)
    print("✅ DEMO COMPLETE!")
    print("=" * 70)


if __name__ == "__main__":
    demo_file_reader()
